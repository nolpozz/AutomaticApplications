"""Render Markdown documents to Markdown / DOCX / PDF.

Markdown is always written. DOCX requires ``python-docx`` and PDF requires
``weasyprint``; when those optional deps are absent we log and skip that format
rather than failing — the pipeline stays runnable on a bare install.
"""

from __future__ import annotations

import re
from pathlib import Path

from job_agent.config.logging import get_logger

logger = get_logger(__name__)


class DocumentRenderer:
    def __init__(self, output_dir: Path | str) -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def render(self, markdown: str, basename: str, formats: list[str]) -> dict[str, str]:
        """Render ``markdown`` into the requested formats. Returns {fmt: path}."""
        paths: dict[str, str] = {}
        base = self.output_dir / basename
        if "md" in formats:
            md_path = base.with_suffix(".md")
            md_path.write_text(markdown, encoding="utf-8")
            paths["md"] = str(md_path)
        if "docx" in formats:
            docx_path = self._render_docx(markdown, base.with_suffix(".docx"))
            if docx_path:
                paths["docx"] = docx_path
        if "pdf" in formats:
            pdf_path = self._render_pdf(markdown, base.with_suffix(".pdf"))
            if pdf_path:
                paths["pdf"] = pdf_path
        return paths

    def _render_docx(self, markdown: str, out: Path) -> str | None:
        try:
            from docx import Document  # python-docx
        except ImportError:
            logger.info("python-docx not installed; skipping DOCX (pip install python-docx)")
            return None
        doc = Document()
        for block in _iter_blocks(markdown):
            if block.kind == "h1":
                doc.add_heading(block.text, level=0)
            elif block.kind == "h2":
                doc.add_heading(block.text, level=1)
            elif block.kind == "h3":
                doc.add_heading(block.text, level=2)
            elif block.kind == "bullet":
                doc.add_paragraph(block.text, style="List Bullet")
            elif block.text.strip():
                doc.add_paragraph(block.text)
        doc.save(str(out))
        return str(out)

    def _render_pdf(self, markdown: str, out: Path) -> str | None:
        try:
            from weasyprint import HTML  # type: ignore
        except (ImportError, OSError):
            logger.info("weasyprint not available; skipping PDF (pip install weasyprint)")
            return None
        HTML(string=_markdown_to_html(markdown)).write_pdf(str(out))
        return str(out)


# -- minimal, dependency-free markdown handling -----------------------------
class _Block:
    def __init__(self, kind: str, text: str) -> None:
        self.kind = kind
        self.text = text


def _iter_blocks(markdown: str) -> list[_Block]:
    blocks: list[_Block] = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            blocks.append(_Block("h3", stripped[4:]))
        elif stripped.startswith("## "):
            blocks.append(_Block("h2", stripped[3:]))
        elif stripped.startswith("# "):
            blocks.append(_Block("h1", stripped[2:]))
        elif stripped.startswith(("- ", "* ")):
            blocks.append(_Block("bullet", _strip_inline(stripped[2:])))
        else:
            blocks.append(_Block("p", _strip_inline(stripped)))
    return blocks


def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def _markdown_to_html(markdown: str) -> str:
    try:
        import markdown as md_lib  # optional, nicer output

        body = md_lib.markdown(markdown)
    except ImportError:
        lines = []
        for block in _iter_blocks(markdown):
            if block.kind.startswith("h"):
                lvl = block.kind[1]
                lines.append(f"<h{lvl}>{block.text}</h{lvl}>")
            elif block.kind == "bullet":
                lines.append(f"<li>{block.text}</li>")
            elif block.text.strip():
                lines.append(f"<p>{block.text}</p>")
        body = "\n".join(lines)
    style = "body{font-family:Georgia,serif;max-width:750px;margin:40px auto;line-height:1.4}"
    return f"<html><head><style>{style}</style></head><body>{body}</body></html>"
